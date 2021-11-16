import sqlite3
import hashlib
import os
import argparse
import time
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

start_time = time.time()

parser = argparse.ArgumentParser(prog='FileChecker', usage='%(prog)s [options]')
parser.add_argument("-H","--hashtype",  choices=["MD5","SHA256"], help="Define o nome da base de dados", default="MD5")
parser.add_argument("-D","--database",  help="Define o nome da base de dados", default="arquivos.db")
parser.add_argument("-P","--path", help="Indica localização dos arquivos", default=".")
parser.add_argument("-L","--limit", help="Número máxima de arquivos no relatório", type=int, default=200)
parser.add_argument("-O","--output", help="Define o nome do arquivo de saida", default="FileReport.docx")
args = parser.parse_args()

DBTABLES = ["ID","Nome","Extensão", "Criado em", "Categoria", "MD5", "SHA256"]

#Conecta ao banco de dados
conn = sqlite3.connect(args.database)

# definindo um cursor
cursor = conn.cursor()

def getInfoFromHash(conn, HASH):
    # lendo os dados
    cursor.execute("""
        SELECT * FROM arquivos WHERE """+args.hashtype+"""= ?
    """, (HASH,))
    return cursor.fetchall()

#Obtém arquivos do diretório/subdiretórios
def getFiles(mypath):
    allFiles = []
    for path, subdirs, files in os.walk(mypath):
        for name in files:
            allFiles.append(os.path.join(path, name))
    return allFiles

#Obtém HASH MD5 de arquivo
def getHashMD5(filename):
    with open(filename, "rb") as f:
        file_hash = hashlib.md5()    
        while byte_block := f.read(8192):
            file_hash.update(byte_block)
    return file_hash.hexdigest()

#Obtém SHA 256 de arquivo
def getHash256(filename):
    with open(filename, "rb") as f:
        file_hash = hashlib.sha256()    
        while byte_block := f.read(8192):
            file_hash.update(byte_block)
    return file_hash.hexdigest()

def genReport(title, files, output):
    files = files[:args.limit]
    document = Document()
    document.add_heading(title, 0)

    style = document.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(20)
    
    head = document.add_heading(f'Arquivos encontrados - {len(files)}', level=1)
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    for file in files:
        if file[0] == None:
            print(file[1], file[2], "Não está na lista")
        else:
            document.add_heading(file[0][1], level=2)
            for ind, table in enumerate(DBTABLES):
                paragraph = document.add_paragraph(table+': ')
                run = paragraph.add_run(f'{file[0][ind]}')
                run.bold = True            
                paragraph.style = 'List Bullet'
            document.add_picture(file[-1], width=Inches(1.5))
    document.save(output)

fileOnPC = []
for file in getFiles(args.path):
    if file != '.\\'+os.path.basename(__file__) and file != '.\\'+args.database:
        fileHash = getHashMD5(file) if args.hashtype else getHash256(file)
        info = getInfoFromHash(conn, fileHash)
        if not info:
            fileOnPC.append([None, file, fileHash])
        else:
            info.append(file)
            fileOnPC.append(info)

genReport("Relatório: Arquivos Ilegais", fileOnPC, args.output)


#Finaliza conexão
conn.close()
